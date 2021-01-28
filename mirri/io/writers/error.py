import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, RGBColor
from datetime import datetime
from inspect import signature
from difflib import SequenceMatcher
from docx2pdf import convert

ERROR_LOG_FOLDER = '.\\logs'
DOCS_FOLDER = '..\\docs' # relative path originating from ERROR_LOG_FOLDER



class Entity():
    """Entity information

    Args:
        acronym: acronym of the entity. Must be a 3-characters captalized string
    """
    def __init__(self, acronym):
        self.entity_acronyms = [func for func in dir(self) if func.isupper() and callable(getattr(self, func)) and not func.startswith("__")]
        self.entity_names = {acr: getattr(self, acr) for acr in self.entity_acronyms}
        self.acronym = acronym
        try:
            self.name = self.entity_names[self.acronym]()
        except KeyError:
            raise ValueError(f'Unknown acronym {self.acronym}.')

    @property
    def name(self):
        return self._name

    @name.setter
    def name(self, name):
        self._name = name

    @property
    def acronym(self):
        return self._acronym
    
    @acronym.setter
    def acronym(self, acronym):
        self._acronym = acronym

    def EFS(self):
        return 'Excel File Structure'

    def GMD(self):
        return 'Growth Media'

    def GOD(self):
        return 'Geographic Origin'

    def LID(self):
        return 'Literature'

    def STD(self):
        return 'Strains'

    def GID(self):
        return 'Genomic Information'

    def UCT(self):
        return 'Uncategorized'


class Error():
    """Error information

    Args
        code_or_message: code or message related to the error. If an error code is specified, param threshold is ignored and the error message is derived \
            from error code. If an error message is specified, an attempt to deduct the error code from the similarity between the specified message and the \
            default messages is performed; If no code is found, a general error code UCT is used instead.

        data: data to be passed to the default messages that require some value. If an error message is specified in code_or_message parameter, this data \
            will be used to help identify the error code.

        threshold: minimum value of similarity between the specified message and the default messages. Should be a value between 0.0 and 1.0. This param \
            is ignoref if a code is specified in code_or_message param.
    """
    def __init__(self, code_or_message, data=None, threshold=0.9):
        self.encoder = self.ErrorMessage()
        self.data = data
        self.threshold = threshold

        if code_or_message in self.encoder.error_codes:
            self.code = code_or_message
            self.message = self.encoder.message(self.code, self.data)
        else:
            self.message = code_or_message
            code = self.find_error_code_v2()
            self.code = code if code != '' else 'UCT'

        self.entity = Entity(self.code[:3])


    def find_error_code_v2(self):
        """Find error code based on the similarity between the default error message and the specified error message

        Returns:
            code (str): error code if any was found or empty string otherwise.
        """
        error = {'code': '', 'ratio': 0.0}
        messages = {code: self.encoder.message(code, self.data) for code in self.encoder.error_codes}
        
        for code, message in messages.items():
            ratio = SequenceMatcher(None, self.message, message).ratio()
            if ratio >= self.threshold and ratio > error['ratio']:
                error['code'] = code
                error['ratio'] = ratio

        return error['code']

    @property
    def entity(self):
        """
            Getter for attribute entity

            return
                entity: entity related to the error (ESF, GMD, GOD, LID, STD, or GID)
        """
        return self._entity

    @entity.setter
    def entity(self, entity):
        """
            Setter for attribute entity

            Args:
                entity: entity related to the error (ESF, GMD, GOD, LID, STD, or GID)
        """
        self._entity = entity

    @property
    def code(self):
        """
            Getter for attribute code

            return
                code: code of the error
        """
        return self._code

    @code.setter
    def code(self, code):
        """
            Setter for attribute code

            Args:
                code: code of the error
        """
        self._code = code

    @property
    def message(self):
        """
            Getter for attribute message

            return
                message: message associated with the error
        """
        return self._message

    @message.setter
    def message(self, message):
        """
            Setter for attribute message

            Args:
                message: message associated with the error
        """
        self._message = message

    @property
    def data(self):
        """
            Getter for attribute data

            return
                data: data used by some error messages. Usually is the primary key of the entry related to the error.
        """
        return self._data

    @data.setter
    def data(self, data):
        """
            Setter for attribute data

            Args:
                data: data used by some error messages. Usually is the primary key of the entry related to the error.
        """
        self._data = data

    # Inner Error Message
    class ErrorMessage():
        """Error messages for each error code"""
        def __init__(self):
            self.error_codes = [func for func in dir(self) if callable(getattr(self, func)) and not func.startswith("__")]
            self.error_messages = {code: getattr(self, code) for code in self.error_codes}

        """
            Get the message associated with the specified error code
        """
        def message(self, code, data=''):
            """get the error message associated with the specified error code
            
            Args:
                code: error code
                data [optional]: a value to be passed to some messages that require some data, for instance, and ID
            """
            if code in self.error_codes:
                sig = signature(self.error_messages[code])
                if len(sig.parameters) > 0:
                    return self.error_messages[code](data)
                else:
                    return self.error_messages[code]()
            else:
                return ''


        """
            Excel File Structure Error Codes
        """
        def EFS01(self):
            return "The 'Growth media' sheet is missing. Please check the provided excel template."

        def EFS02(self):
            return "The 'Geographic origin' sheet is missing. Please check the provided excel template."

        def EFS03(self):
            return "The 'Literature' sheet is missing. Please check the provided excel template."

        def EFS04(self):
            return "The 'Sexual state' sheet is missing. Please check the provided excel template."

        def EFS05(self):
            return "The 'Strains' sheet is missing. Please check the provided excel template."

        def EFS06(self):
            return "The 'Ontobiotope' sheet is missing. Please check the provided excel template."

        def EFS07(self):
            return "The 'Markers' sheet is missing. Please check the provided excel template."

        def EFS08(self):
            return "The 'Genomic information' sheet is missing. Please check the provided excel template."


        """
            Growth Media Error Codes
        """
        def GMD01(self):
            return "The 'Acronym' is a mandatory field. The column can not be empty."

        def GMD02(self):
            return "The 'Description' is a mandatory field. The column can not be empty."


        """
            Geographic Origin Error Codes
        """
        def GOD01(self):
            return "The 'ID' is a mandatory field. The column can not be empty."

        def GOD02(self):
            return "The 'Country' is a mandatory field. The column can not be empty."

        def GOD03(self, name):
            return f"The 'Country' named as {name} is incorrect."

        def GOD04(self, country):
            return f"The 'Locality' is missing for Country {country}."

        """
            Literature Error Codes
        """
        def LID01(self):
            return "The 'ID' is a mandatory field. The column can not be empty."

        def LID02(self):
            return "The 'Full reference' is a mandatory field. The column can not be empty."

        def LID03(self):
            return "The 'Authors' is a mandatory field. The column can not be empty."

        def LID04(self, id):
            return f"The 'Authors' is missing for ID number {id}."

        def LID05(self):
            return "The 'Title' is a mandatory field. The column can not be empty."

        def LID06(self, id):
            return f"The 'Title' is missing for ID number {id}."

        def LID07(self):
            return "The 'Journal' is a mandatory field. The column can not be empty."

        def LID08(self, id):
            return f"The 'Journal' is missing for ID number {id}."

        def LID09(self):
            return "The 'Year' is a mandatory field. The column can not be empty."

        def LID10(self, id):
            return f"The 'Year' is missing for ID number {id}."

        def LID11(self):
            return "The 'Volume' is a mandatory field. The column can not be empty."

        def LID12(self, id):
            return f"The 'Volume' is missing for ID number {id}."

        def LID13(self):
            return "The 'First page' is a mandatory field. The column can not be empty."

        def LID14(self, id):
            return f"The 'First page' is missing for ID number {id}."

        """
            Strains Error Codes
        """
        def STD01(self):
            return "The 'Accession number' is a mandatory field. The column can not be empty."

        def STD02(self):
            return "The 'Accession number' is not according to the specification."

        def STD03(self, accession_number):
            return f"The 'Other culture collection numbers' for strain with Accession Number {accession_number} is not according to the specification."

        def STD04(self):
            return "The 'Restriction on use' is a mandatory field. The column can not be empty."

        def STD05(self, accession_number):
            return f"The 'Restriction on use' for strain with Accession Number {accession_number} is not according to the specification."

        def STD06(self):
            return "The 'Nagoya protocol restrictions and compliance conditions' is a mandatory field. The column can not be empty."

        def STD07(self, accession_number):
            return f"The 'Nagoya protocol restrictions and compliance conditions' for strain with Accession Number {accession_number} is not according to the specification."

        def STD08(self, accession_number):
            return f"The 'ABS related files' for strain with Accession Number {accession_number} is not a valid URL."

        def STD09(self, accession_number):
            return f"The 'MTA file' for strain with Accession Number {accession_number} is not a valid URL."

        def STD10(self, accession_number):
            return f"The 'Strain from a registered collection' for strain with Accession Number {accession_number} is not according to specification."

        def STD11(self):
            return "The 'Risk group' is a mandatory field. The column can not be empty."

        def STD12(self, accession_number):
            return f"The 'Risk group' for strain with Accession Number {accession_number} is not according to specification."

        def STD13(self, accession_number):
            return f"The 'Dual use' for strain with Accession Number {accession_number} is not according to specification."

        def STD14(self, accession_number):
            return f"The “Quarentine in europe” entered for strain with Accession Number {accession_number} is incorrect."

        def STD15(self):
            return "The 'Organism type' is a mandatory field. The column can not be empty."

        def STD16(self, accession_number):
            return f"The 'Organism type' for strain with Accession Number {accession_number} is incorrect."

        def STD17(self):
            return "The 'Taxon name' is a mandatory field. The column can not be empty."

        def STD18(self, accession_number):
            return f"The 'Taxon name' is missing for strain with Accession Number {accession_number}."

        def STD19(self, accession_number):
            return f"The 'Infrasubspecific names' for strain with Accession Number {accession_number} is incorrect."

        def STD20(self, accession_number):
            return f"The 'History of deposit' sequence for strain with Accession Number {accession_number} is incorrect."

        def STD21(self, accession_number):
            return f"The 'Date of deposit' for strain with Accession Number {accession_number} is incorrect."

        def STD22(self, accession_number):
            return f"The 'Date of collection' for strain with Accession Number {accession_number} is incorrect."

        def STD23(self, accession_number):
            return f"The 'Date of isolation' for strain with Accession Number {accession_number} is incorrect."

        def STD24(self, accession_number):
            return f"The 'Date of inclusion in the catalogue' for strain with Accession Number {accession_number} is incorrect."

        def STD25(self, accession_number):
            return f"The 'Tested temperature growth range' for strain with Accession Number {accession_number} is incorrect."

        def STD26(self):
            return f"The 'Recommended growth temperature' is a mandatory field for eash strain. The column can not be empty."

        def STD27(self, accession_number):
            return f"The 'Recommended growth temperature' is missing for strain with Accession Number {accession_number}."

        def STD28(self):
            return "The 'Recommended medium for growth' is a mandatory field. The column can not be empty."

        def STD29(self, accession_number):
            return f"The 'Recommended medium for growth' is missing for strain with Accession Number {accession_number}."

        def STD30(self):
            return "The 'Forms of supply' is a mandatory field. The column can not be empty."

        def STD31(self, accession_number):
            return f"The 'Forms of supply' is missing for strain with Accession Number {accession_number}."

        def STD32(self, accession_number):
            return f"The 'Coordinates of geographic origin' for strain with Accession Number {accession_number} are incorrect."

        def STD33(self, accession_number):
            return f"The 'Altitude of geographic origin' for strain with Accession Number {accession_number} are incorrect."

        def STD34(self):
            return "The 'Geographic origin' is a mandatory field. The column can not be empty."

        def STD35(self, accession_number):
            return f"The 'Geographic origin' is missing for strain with Accession Number {accession_number}."

        def STD36(self, accession_number):
            return f"The 'GMO' for strain with Accession Number {accession_number} is incorrect."

        def STD37(self, accession_number):
            return f"The 'Literature' for strain with Accession Number {accession_number} is incorrect."

        def STD38(self, accession_number):
            return f"The 'Sexual state' for strain with Accession Number {accession_number} is incorrect."

        def STD39(self, accession_number):
            return f"The 'Ploidy' for strain with Accession Number {accession_number} is incorrect."

        def STD40(self, accession_number):
            return f"The 'Interspecific hybrid' for strain with Accession Number {accession_number} is incorrect."

        def STD41(self, accession_number):
            return f"The 'Plasmids collection fields' for strain with Accession Number {accession_number} is incorrect."

        def STD42(self, accession_number):
            return f"The 'Ontobiotope term for the isolation habitat' for strain with Accession Number {accession_number} is incorrect."

        def STD43(self, accession_number):
            return f"The 'Literature linked to the sequence/genome' for strain with Accession Number {accession_number} is incorrect."

        """
            Genomic Information Error Codes
        """
        def GID01(self, accession_number):
            return f"The 'Strain Acession Number' (Strain AN) with ID {accession_number} is incorrect."

        def GID02(self, accession_number):
            return f"The 'Marker' for Strain with ID {accession_number} is incorrect."

        def GID03(self, accession_number):
            return f"The 'INSDC AN' for Strain with ID {accession_number} is incorrect."

        def GID04(self, accession_number):
            return f"The 'Sequence' for Strain with ID {accession_number} is incorrect."


class ErrorLog():
    """
        Error Logging

        Args:
            input_filename: name of the file which the error log is derived from
            cc [optional]: culture collection identifier
            date [optional]: date the inputed file was submited for validation (date of last modification)
    """
    def __init__(self, input_filename: str, cc: str=None, date: str = None):
        self.input_filename = input_filename
        self.cc = cc
        self.date = datetime.strptime(date, '%d-%m-%Y').date() if date is not None else None
        self.id = 0
        self.errors = {}
        self.document = docx.Document(f'.\\docs\\Error_Log_Style_Sheet.docx')
            

    def write(self, path: str):
        """Write erros to log file

        Args:
            path (str): path of the file to write the errors log
        """
        def hyperlink(paragraph, text, url):
            """Generate a hyperlink text

            Args:
                paragraph (Paragraph): Paragraph object to append the hyperlink
                text (str): text of the hyperlink
                url (str): the url to which the hyperlink points to

            Returns:
                hyperlink: hyperlink object
            """
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')

            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)

            r = paragraph.add_run()
            r._r.append(hyperlink)

            r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
            r.font.underline = True

            return hyperlink

        heading = self.document.add_heading('Error Log', 0)
        heading.style = self.document.styles['Title']

        cc = f' of Culture Collection {self.cc}' if self.cc is not None else ''
        date = f'in {self.date} ' if self.date is not None else ''
        mail_to = 'ict-support@mirri.org'
        subject = 'Validator Error Log'

        self.document.add_paragraph(f'Dear Curator{cc},')
        paragraph = self.document.add_paragraph(f'the Excel File {self.input_filename} you\'ve provided {date}with your collection Strains Data contains errors/missing data. ')
        paragraph.add_run('Please, see below the list of detected errors/missing data, for you to proceed with the appropriated correction/completion.')

        paragraph = self.document.add_paragraph('If you need help, please refer to the instructions contained in "')
        hyperlink(paragraph, 'ICT-TaskForce_HowToCompileTheSheets_v20200601.pdf', f'{DOCS_FOLDER}\\ICT-TaskForce_HowToCompileTheSheets_v20200601.pdf')
        paragraph.add_run('" and "')
        hyperlink(paragraph, 'ICT-TaskForce_RecommendationsToCollections_v20200601.pdf', f'{DOCS_FOLDER}\\ICT-TaskForce_RecommendationsToCollections_v20200601.pdf')
        paragraph.add_run('".\nYou can also contact the MIRRI ICT by email using ')
        hyperlink(paragraph, 'ICT Support', f'mailto:{mail_to}?Subject={subject}')

        self.document.add_page_break()
        
        self.document.add_heading(f'Analysis of {Entity("EFS").name}', level=1).style = self.document.styles['Heading 1']
        self.document.add_paragraph('The structure of your Excel File show the following changes, as compared to the original Template:')

        table = self.document.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells = hdr_cells[0].merge(hdr_cells[1])
        hdr_cells.text = Entity('EFS').name
        hdr_cells.paragraphs[0].style = self.document.styles['Table Header']
        subhdr_cells = table.rows[1].cells
        subhdr_cells[0].text = 'Error Code'
        subhdr_cells[0].paragraphs[0].style = self.document.styles['Table Header']
        subhdr_cells[0].width = Cm(4.0)
        subhdr_cells[1].text = 'Error Message'
        subhdr_cells[1].paragraphs[0].style = self.document.styles['Table Header']

        if 'EFS' in self.errors:
            for error in sorted(self.errors['EFS'], key=lambda e: e.code):
                row_cells = table.add_row().cells
                row_cells[0].text = error.code
                row_cells[0].paragraphs[0].style = self.document.styles['Table Cell']
                row_cells[0].width = Cm(4.0)
                row_cells[1].text = error.message
                row_cells[1].paragraphs[0].style = self.document.styles['Table Cell']

        self.document.add_page_break()


        self.document.add_heading('Analysis of Data Set', level=1).style = self.document.styles['Heading 1']
        self.document.add_paragraph('Your Data shows the following errors or missing items:')

        for entity_acronym in self.errors:
            if entity_acronym in ['EFS', 'UCT']: continue
            
            entity = Entity(entity_acronym)
            self.document.add_heading(entity.name, level=2).style = self.document.styles['Heading 2']
            self.document.add_paragraph(f'The “{entity.name}” Sheet in your Excel File shows the following errors or missing items:')

            table = self.document.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells = hdr_cells[0].merge(hdr_cells[1])
            hdr_cells.text = entity.name
            hdr_cells.paragraphs[0].style = self.document.styles['Table Header']
            subhdr_cells = table.rows[1].cells
            subhdr_cells[0].text = 'Error Code'
            subhdr_cells[0].paragraphs[0].style = self.document.styles['Table Header']
            subhdr_cells[0].width = Cm(4.0)
            subhdr_cells[1].text = 'Error Message'
            subhdr_cells[1].paragraphs[0].style = self.document.styles['Table Header']

            if entity.acronym in self.errors:
                for error in sorted(self.errors[entity.acronym], key=lambda e: e.code):
                    row_cells = table.add_row().cells
                    row_cells[0].text = error.code
                    row_cells[0].paragraphs[0].style = self.document.styles['Table Cell']
                    row_cells[0].width = Cm(4.0)
                    row_cells[1].text = error.message
                    row_cells[1].paragraphs[0].style = self.document.styles['Table Cell']
                
        if 'UCT' in self.errors:
            self.document.add_page_break()
            self.document.add_heading('Uncategorized Errors', level=1).style = self.document.styles['Heading 1']
            self.document.add_paragraph('The following errors were also identified while validating your data:')
            
            entity = Entity('UCT')

            table = self.document.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells = hdr_cells[0].merge(hdr_cells[1])
            hdr_cells.text = entity.name
            hdr_cells.paragraphs[0].style = self.document.styles['Table Header']
            subhdr_cells = table.rows[1].cells
            subhdr_cells[0].text = 'Error Code'
            subhdr_cells[0].paragraphs[0].style = self.document.styles['Table Header']
            subhdr_cells[0].width = Cm(4.0)
            subhdr_cells[1].text = 'Error Message'
            subhdr_cells[1].paragraphs[0].style = self.document.styles['Table Header']
            
            for error in self.errors['UCT']:
                row_cells = table.add_row().cells
                row_cells[0].text = error.code
                row_cells[0].paragraphs[0].style = self.document.styles['Table Cell']
                row_cells[0].width = Cm(4.0)
                row_cells[1].text = error.message
                row_cells[1].paragraphs[0].style = self.document.styles['Table Cell']


        try:
            self.document.save(f'{path}\\{self.input_filename}_error_log.docx')
            convert(f'{path}\\{self.input_filename}_error_log.docx')
        except:
            raise



    def __str__(self):
        limit = 200
        count = 0
        to_print = f'Error log for file <{self.input_filename}> sent by <{self.cc}> on <{self.date}>\n'
        to_print += f'printing first {limit} errors\n\n'
        to_print += '{:<5} | {:<10} | {:<10} | {:<250}\n'.format('#', 'ENTITY', 'CODE', 'MESSAGE')
        for _, errors in self.errors.items():
            if count == limit:
                    break

            for error in errors:
                count += 1
                if count == limit:
                    break
                to_print += f'{count:>05}'
                to_print += f' | {error.entity.acronym:<10}'
                to_print += f' | {error.code:<10}'
                to_print += f' | {error.message[:248]:<250}...' if len(error.message) > 250 else f' | {error.message:<250}'
                to_print += '\n'

        return to_print


    @property
    def input_filename(self):
        """
            Getter for input filename

            return
                input filename [str]: name of the file which the error log is derived from
        """
        return self._input_filename


    @input_filename.setter
    def input_filename(self, input_filename: str):
        """
            Setter for input filename

            Args:
                input filename (str): name of the file which the error log is derived from
        """
        self._input_filename = input_filename


    @property
    def cc(self):
        """
            Getter for culture collection identifier

            return
                cc [str]: culture collection identifier
        """
        return self._cc


    @cc.setter
    def cc(self, cc: str):
        """
            Setter for culture collection identifier

            Args:
                cc (str): culture collection identifier
        """
        self._cc = cc


    @property
    def date(self):
        """
            Getter for date the inputed file was submited for validation

            return
                date [str]: date the inputed file was submited for validation
        """
        return self._date


    @date.setter
    def date(self, date):
        """
            Setter for date the inputed file was submited for validation

            Args:
                date (str): date the inputed file was submited for validation
        """
        self._date = date


    def get_errors(self):
        """
            Getter for errors identified

            return
                errors [dict]: errors identified
        """
        return self.errors


    def add_error(self, error: Error):
        """
            Add an error

            Args:
                error (Error): error to be added
        """
        if error.entity.acronym not in self.errors:
            self.errors[error.entity.acronym] = [error]
        else:
            self.errors[error.entity.acronym].append(error)


if __name__ == '__main__':
    error_log = ErrorLog('MIRRI-IS_dataset_BEA_template_30092020', 'BEA', '30-09-2020')
    errors = [
        Error('EFS01'),
        Error('EFS02'),
        Error('GOD01'),
        Error('GOD04', 'Brazil'),
        Error('LID01'),
        Error('STD02'),
        Error('STD07', 'AN 123'),
        Error('STD16'),
        Error('STD23', 'AN 123'),
        Error('GID01', 'AN 456'),
        Error('GID03', 'AN 456'),
        Error('The “Acronym” is a mandatory field. The Column can not be empty.'),
        Error('The “Country” named as Spain is incorrect', 'Spain'),
        Error('The Growth Medium ABC 123 is not in the Growth Media datasheet.')
    ]
    
    for error in errors:
        error_log.add_error(error)

    print(error_log)
    # print(f'Writing to file \'{ERROR_LOG_FOLDER}\Error_Log_Example.docx\'')
    # error_log.write(ERROR_LOG_FOLDER)