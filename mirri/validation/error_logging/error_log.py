from typing import Optional
import docx
import os
import mirri
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Cm
from datetime import datetime
from docx2pdf import convert
from tempfile import NamedTemporaryFile
from .error import Error, Entity

DOCS_FOLDER = os.path.abspath(os.path.join(
    os.path.dirname(mirri.__file__), '..', 'docs'))


class ErrorLog():
    """
        Error Logging

        Args:
            input_filename: name of the file which the error log is derived from
            cc [optional]: culture collection identifier
            date [optional]: date the inputed file was submited for validation (date of last modification)
            limit [int]: maximum number fo errors to be writen on the output file
    """

    def __init__(self, input_filename: str, cc: str = None, date: str = None, limit: int = 100):
        self.input_filename = input_filename
        self.cc = cc
        self.date = datetime.strptime(
            date, r'%d-%m-%Y').date() if date is not None else None

        self._errors = {}
        self.limit = limit
        self._counter = 0
        self.document = docx.Document(self._fpath_to_style_doc)

    def write(self, path: str):
        """Write erros to log file

        Args:
            path (str): path of the file to write the errors log

        Return:
            (str) path to the PDF file
        """
        heading = self.document.add_heading('Error Log', 0)
        heading.style = self.document.styles['Title']

        cc = f' of Culture Collection {self.cc}' if self.cc is not None else ''
        date = f'in {self.date} ' if self.date is not None else ''
        mail_to = 'ict-support@mirri.org'
        subject = 'Validator Error Log'

        self.document.add_paragraph(f'Dear Curator{cc},')
        paragraph = self.document.add_paragraph(
            f'the Excel File {self.input_filename} you\'ve provided {date}with your collection Strains Data contains errors/missing data. ')
        paragraph.add_run(
            'Please, see below the list of detected errors/missing data, for you to proceed with the appropriated correction/completion.')

        paragraph = self.document.add_paragraph(
            'If you need help, please refer to the instructions contained in "')
        self._hyperlink(
            paragraph, 'ICT-TaskForce_HowToCompileTheSheets_v20200601.pdf', self._fpath_how_to_compile)
        paragraph.add_run('" and "')
        self._hyperlink(
            paragraph, 'ICT-TaskForce_RecommendationsToCollections_v20200601.pdf', self._fpath_recommendations)
        paragraph.add_run(
            '".\nYou can also contact the MIRRI ICT by email using ')
        self._hyperlink(paragraph, 'ICT Support',
                        f'mailto:{mail_to}?Subject={subject}')

        if 'EFS' in self._errors:
            entity = Entity("EFS")
            self.document.add_page_break()
            self.document.add_heading(
                f'Analysis of {entity.name}', level=1).style = self.document.styles['Heading 1']
            self.document.add_paragraph(
                'The structure of your Excel File show the following changes, as compared to the original Template:')
            self._error_table(entity)

        if self._counter == self.limit:
            self._limit_message()
        elif len(self._errors.keys()) > 1:
            self.document.add_page_break()
            self.document.add_heading(
                'Analysis of Data Set', level=1).style = self.document.styles['Heading 1']
            self.document.add_paragraph(
                'Your Data shows the following errors or missing items:')

            for entity_acronym in self._errors:
                if entity_acronym in ['EFS', 'UCT']:
                    continue
                entity = Entity(entity_acronym)
                self.document.add_heading(
                    entity.name, level=2).style = self.document.styles['Heading 2']
                self.document.add_paragraph(
                    f'The “{entity.name}” Sheet in your Excel File shows the following errors or missing items:')
                self._error_table(entity, lambda e: (
                    e.data is not None, e.data))
                if self._counter == self.limit:
                    break

            if self._counter == self.limit:
                self._limit_message()
            elif 'UCT' in self._errors:
                self.document.add_page_break()
                self.document.add_heading(
                    'Uncategorized Errors', level=1).style = self.document.styles['Heading 1']
                self.document.add_paragraph(
                    'The following errors were also identified while validating your data:')
                self._error_table(Entity('UCT'), lambda e: (
                    e.data is not None, e.data))
                if self._counter == self.limit:
                    self._limit_message()

        try:
            docx_fhand = NamedTemporaryFile(
                dir=path, suffix='_error_log.docx', delete=False)
            self.document.save(docx_fhand)
            docx_fhand.close()
            pdf_fhand = os.path.join(
                path, f'{self.input_filename}_error_log.pdf')
            convert(docx_fhand.name, pdf_fhand)
        finally:
            if not docx_fhand.closed:
                docx_fhand.close()
            os.unlink(docx_fhand.name)

        if os.path.exists(pdf_fhand):
            return pdf_fhand
        else:
            return None

    def _hyperlink(self, paragraph, text, url):
        """Generate a hyperlink text

        Args:
            paragraph (Paragraph): Paragraph object to append the hyperlink
            text (str): text of the hyperlink
            url (str): the url to which the hyperlink points to

        Returns:
            hyperlink: hyperlink object
        """
        part = paragraph.part
        r_id = part.relate_to(
            url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        r = paragraph.add_run()
        r._r.append(hyperlink)

        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return hyperlink

    def _error_table(self, entity, sort_func=None):
        table = self.document.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells = hdr_cells[0].merge(hdr_cells[1])
        hdr_cells.text = entity.name
        hdr_cells.paragraphs[0].style = self.document.styles['Table Header']
        subhdr_cells = table.rows[1].cells
        subhdr_cells[0].text = 'Identifier'
        subhdr_cells[0].paragraphs[0].style = self.document.styles['Table Header']
        subhdr_cells[0].width = Cm(4.0)
        subhdr_cells[1].text = 'Error Message'
        subhdr_cells[1].paragraphs[0].style = self.document.styles['Table Header']

        if sort_func is not None:
            errors = sorted(self._errors[entity.acronym], key=sort_func)
        else:
            errors = self._errors[entity.acronym]

        for error in errors:
            row_cells = table.add_row().cells
            row_cells[0].text = error.data
            row_cells[0].paragraphs[0].style = self.document.styles['Table Cell']
            row_cells[0].width = Cm(4.0)
            row_cells[1].text = error.message
            row_cells[1].paragraphs[0].style = self.document.styles['Table Cell']
            self._counter += 1
            if self._counter == self.limit:
                break

        return table

    def _limit_message(self):
        self.document.add_page_break()
        self.document.add_paragraph(
            'Your file contains too many erros and therefore this document was truncated. Please resolve the aforementioned errors and resubmit the excel file.'
        )

    @property
    def _fpath_to_style_doc(self) -> str:
        return os.path.join(DOCS_FOLDER, 'Error_Log_Style_Sheet.docx')

    @property
    def _fpath_how_to_compile(self) -> str:
        return os.path.join(
            DOCS_FOLDER,
            'ICT-TaskForce_HowToCompileTheSheets_v20200601.pdf'
        )

    @property
    def _fpath_recommendations(self) -> str:
        return os.path.join(
            DOCS_FOLDER,
            'ICT-TaskForce_RecommendationsToCollections_v20200601.pdf'
        )

    @property
    def input_filename(self) -> str:
        return self._input_filename

    @input_filename.setter
    def input_filename(self, input_filename: str) -> None:
        self._input_filename = input_filename

    @property
    def cc(self) -> str:
        return self._cc

    @cc.setter
    def cc(self, cc: str) -> None:
        self._cc = cc

    @property
    def date(self) -> Optional[datetime]:
        return self._date

    @date.setter
    def date(self, date: Optional[datetime] = None) -> None:
        self._date = date

    def get_errors(self) -> dict:
        return self._errors

    def add_error(self, error: Error) -> list:
        if error.entity.acronym not in self._errors:
            self._errors[error.entity.acronym] = [error]
        else:
            self._errors[error.entity.acronym].append(error)
